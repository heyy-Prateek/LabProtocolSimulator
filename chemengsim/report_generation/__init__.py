"""
Report Generation Module for Chemical Engineering Simulator
==========================================================

This module provides functions to generate reports from experiment data.
"""

import os
import time
import datetime
from typing import Dict, Any, Optional, List

def generate_report(
    experiment_id: str,
    title: Optional[str] = None,
    description: Optional[str] = None,
    include_charts: bool = True,
    format: str = "docx",
    data: Optional[Dict[str, Any]] = None
) -> Optional[str]:
    """
    Generate a report for an experiment.
    
    Args:
        experiment_id: ID of the experiment
        title: Report title (optional)
        description: Report description (optional)
        include_charts: Whether to include charts in the report
        format: Report format ('docx' or 'pdf')
        data: Experiment data (optional, will use current data if None)
        
    Returns:
        Path to the generated report file, or None if report generation failed
    """
    try:
        # Import necessary modules
        from docx import Document
        
        # Import the experiment module
        from experiments import get_experiment
        
        # Get experiment module
        experiment = get_experiment(experiment_id)
        if not experiment:
            print(f"Error: Experiment with ID '{experiment_id}' not found")
            return None
            
        # Create report title
        if not title:
            title = f"{experiment.name} Report"
            
        # Create description
        if not description:
            description = f"Report generated for {experiment.name} experiment on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading(title, 0)
        
        # Add description
        doc.add_paragraph(description)
        
        # Add experiment details
        doc.add_heading("Experiment Details", level=1)
        doc.add_paragraph(f"Name: {experiment.name}")
        doc.add_paragraph(f"Description: {experiment.description}")
        
        # Add parameters
        doc.add_heading("Parameters", level=1)
        
        # If we don't have data, use default parameters
        parameters = data.get("parameters", experiment.default_parameters) if data else experiment.default_parameters
        
        # Add parameter table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Parameter"
        header_cells[1].text = "Value"
        
        # Add parameters
        for name, value in parameters.items():
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = str(value)
            
        # Add equations
        if hasattr(experiment, "equations") and experiment.equations:
            doc.add_heading("Equations", level=1)
            for equation in experiment.equations:
                doc.add_paragraph(equation)
                
        # Add variables
        if hasattr(experiment, "variables") and experiment.variables:
            doc.add_heading("Variables", level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Variable"
            header_cells[1].text = "Description"
            
            # Add variables
            for name, description in experiment.variables.items():
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = description
        
        # Add results
        if data and "results" in data:
            doc.add_heading("Results", level=1)
            
            # Add results table
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Metric"
            header_cells[1].text = "Value"
            
            # Add results
            for name, value in data["results"].items():
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = str(value)
                
        # Add charts
        if include_charts and data and "charts" in data:
            doc.add_heading("Charts", level=1)
            
            # For now, just add placeholder text for charts
            doc.add_paragraph("Charts would be included here in a real implementation.")
            
        # Create reports directory if it doesn't exist
        reports_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), "reports")
        os.makedirs(reports_dir, exist_ok=True)
        
        # Save report
        timestamp = int(time.time())
        filename = f"{experiment_id}_{timestamp}.{format}"
        filepath = os.path.join(reports_dir, filename)
        
        doc.save(filepath)
        
        print(f"Report generated: {filepath}")
        return filepath
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        return None

def get_report_templates() -> List[Dict[str, Any]]:
    """
    Get available report templates.
    
    Returns:
        List of report template information
    """
    # For now, just return a simple list of templates
    return [
        {
            "id": "default",
            "name": "Default Template",
            "description": "Standard report template with experiment details and results."
        },
        {
            "id": "detailed",
            "name": "Detailed Template",
            "description": "Comprehensive report template with detailed analysis and extended charts."
        },
        {
            "id": "summary",
            "name": "Summary Template",
            "description": "Brief summary report with key results only."
        }
    ]

def get_report_formats() -> List[Dict[str, Any]]:
    """
    Get available report formats.
    
    Returns:
        List of report format information
    """
    return [
        {
            "id": "docx",
            "name": "Word Document (.docx)",
            "description": "Microsoft Word document format."
        },
        {
            "id": "pdf",
            "name": "PDF Document (.pdf)",
            "description": "Portable Document Format."
        }
    ]